import tkinter as tk
from tkinter import ttk, messagebox
import ttkbootstrap as tb
from ttkbootstrap.constants import *
import mysql.connector
from mysql.connector import Error
from docx import Document
import os
from datetime import datetime
import tempfile
import re
import pythoncom
import subprocess
from docx2pdf import convert
import sys

class EmployeeFormApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Employee Data Management")
        self.root.geometry("1000x700")
        
        # MySQL connection details
        self.db_config = {
            'host': 'localhost',
            'database': 'employee_db',
            'user': 'root',
            'password': 'admin123'  # Change to your MySQL password
        }
        
        # Store selected record
        self.selected_record = None
        
        # Template path - using your existing template
        self.template_path = r"C:\Users\Rohtash\AppData\Local\Temp\EmployeeTemplates\employee_template.docx"
        
        self.setup_database()
        self.create_widgets()
        self.load_data()
    
    def setup_database(self):
        """Create database and table if they don't exist"""
        try:
            # First connect without database to create it
            conn = mysql.connector.connect(
                host=self.db_config['host'],
                user=self.db_config['user'],
                password=self.db_config['password']
            )
            cursor = conn.cursor()
            
            # Create database if not exists
            cursor.execute("CREATE DATABASE IF NOT EXISTS employee_db")
            conn.commit()
            
            # Close initial connection
            cursor.close()
            conn.close()
            
            # Now connect to the specific database
            conn = mysql.connector.connect(**self.db_config)
            cursor = conn.cursor()
            
            # Create table if not exists
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS Table1 (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    name VARCHAR(255) NOT NULL,
                    age INT NOT NULL,
                    salary DECIMAL(10, 2) NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.commit()
            
            cursor.close()
            conn.close()
            
        except Error as e:
            messagebox.showerror("Database Error", f"Error setting up database: {e}")
    
    def create_connection(self):
        """Create database connection"""
        try:
            conn = mysql.connector.connect(**self.db_config)
            return conn
        except Error as e:
            messagebox.showerror("Database Error", f"Error connecting to database: {e}")
            return None
    
    def format_currency(self, amount):
        """Format amount as Indian Rupees with proper formatting"""
        try:
            # Remove any existing currency symbols and commas
            clean_amount = re.sub(r'[^\d.]', '', str(amount))
            numeric_amount = float(clean_amount)
            
            # Format with Indian numbering system (lakhs and crores)
            if numeric_amount >= 10000000:  # 1 crore
                formatted = f"₹{numeric_amount/10000000:.2f} crore"
            elif numeric_amount >= 100000:  # 1 lakh
                formatted = f"₹{numeric_amount/100000:.2f} lakh"
            else:
                # Format with commas for thousands
                formatted = f"₹{numeric_amount:,.2f}"
            
            return formatted
        except ValueError:
            return f"₹{amount}"
    
    def parse_currency(self, currency_str):
        """Parse currency string to float, handling Indian Rupees format"""
        try:
            # Remove ₹ symbol and any text like 'lakh', 'crore'
            clean_str = re.sub(r'[₹\$,]', '', str(currency_str))
            clean_str = re.sub(r'[a-zA-Z\s]', '', clean_str)
            return float(clean_str)
        except ValueError:
            raise ValueError("Invalid salary format. Please enter a valid number.")
    
    def create_widgets(self):
        # Main frame
        main_frame = tb.Frame(self.root, padding=20)
        main_frame.pack(fill=BOTH, expand=True)
        
        # Title
        title_label = tb.Label(main_frame, text="Employee Information System", 
                              font=("Helvetica", 18, "bold"), bootstyle=PRIMARY)
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20))
        
        # Form frame
        form_frame = tb.LabelFrame(main_frame, text="Employee Details", padding=15)
        form_frame.grid(row=1, column=0, sticky=NW, padx=(0, 20))
        
        # Name field
        tb.Label(form_frame, text="Name:").grid(row=0, column=0, sticky=W, pady=5)
        self.name_var = tk.StringVar()
        self.name_entry = tb.Entry(form_frame, textvariable=self.name_var, width=30)
        self.name_entry.grid(row=0, column=1, pady=5, padx=(10, 0))
        
        # Age field
        tb.Label(form_frame, text="Age:").grid(row=1, column=0, sticky=W, pady=5)
        self.age_var = tk.StringVar()
        self.age_entry = tb.Entry(form_frame, textvariable=self.age_var, width=30)
        self.age_entry.grid(row=1, column=1, pady=5, padx=(10, 0))
        
        # Salary field with example placeholder
        tb.Label(form_frame, text="Salary (₹):").grid(row=2, column=0, sticky=W, pady=5)
        self.salary_var = tk.StringVar()
        self.salary_entry = tb.Entry(form_frame, textvariable=self.salary_var, width=30)
        self.salary_entry.grid(row=2, column=1, pady=5, padx=(10, 0))
        # Add placeholder text
        self.salary_entry.insert(0, "e.g., 50000 or 5.5 lakh")
        self.salary_entry.config(foreground="gray")
        self.salary_entry.bind("<FocusIn>", self.clear_placeholder)
        self.salary_entry.bind("<FocusOut>", self.add_placeholder)
        
        # Button frame
        button_frame = tb.Frame(form_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=20)
        
        # Submit button
        submit_btn = tb.Button(button_frame, text="Submit", 
                              command=self.submit_data, 
                              bootstyle=SUCCESS, width=12)
        submit_btn.pack(side=LEFT, padx=(0, 10))
        
        # Print button
        self.print_btn = tb.Button(button_frame, text="Generate PDF", 
                                  command=self.generate_pdf, 
                                  bootstyle=INFO, width=12, state=DISABLED)
        self.print_btn.pack(side=LEFT)
        
        # Clear selection button
        clear_btn = tb.Button(button_frame, text="Clear Selection", 
                             command=self.clear_selection, 
                             bootstyle=WARNING, width=12)
        clear_btn.pack(side=LEFT, padx=(10, 0))
        
        # Table frame
        table_frame = tb.LabelFrame(main_frame, text="Employee Records", padding=10)
        table_frame.grid(row=1, column=1, sticky=NSEW)
        
        # Create treeview (table)
        columns = ("ID", "Name", "Age", "Salary", "Created At")
        self.tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=15)
        
        # Define headings
        self.tree.heading("ID", text="ID")
        self.tree.heading("Name", text="Name")
        self.tree.heading("Age", text="Age")
        self.tree.heading("Salary", text="Salary")
        self.tree.heading("Created At", text="Created At")
        
        # Define column widths
        self.tree.column("ID", width=50, anchor=CENTER)
        self.tree.column("Name", width=150, anchor=W)
        self.tree.column("Age", width=80, anchor=CENTER)
        self.tree.column("Salary", width=120, anchor=E)
        self.tree.column("Created At", width=120, anchor=CENTER)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        # Grid layout
        self.tree.grid(row=0, column=0, sticky=NSEW)
        scrollbar.grid(row=0, column=1, sticky=NS)
        
        # Bind selection event
        self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)
        
        # Status label
        self.status_label = tb.Label(main_frame, text="No record selected", bootstyle=SECONDARY)
        self.status_label.grid(row=2, column=0, columnspan=2, pady=(20, 0), sticky=W)
        
        # Template info
        template_info = tb.Label(main_frame, text=f"Template: {self.template_path}", 
                                bootstyle=SECONDARY, font=("Helvetica", 8))
        template_info.grid(row=3, column=0, columnspan=2, pady=(5, 0), sticky=W)
        
        # Configure grid weights
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)
        table_frame.columnconfigure(0, weight=1)
        table_frame.rowconfigure(0, weight=1)
    
    def clear_placeholder(self, event):
        """Clear placeholder text when entry gets focus"""
        if self.salary_entry.get() == "e.g., 50000 or 5.5 lakh":
            self.salary_entry.delete(0, tk.END)
            self.salary_entry.config(foreground="black")
    
    def add_placeholder(self, event):
        """Add placeholder text if entry is empty"""
        if not self.salary_entry.get():
            self.salary_entry.insert(0, "e.g., 50000 or 5.5 lakh")
            self.salary_entry.config(foreground="gray")
    
    def on_tree_select(self, event):
        """Handle treeview selection"""
        selected_items = self.tree.selection()
        if selected_items:
            item = selected_items[0]
            self.selected_record = self.tree.item(item, 'values')
            
            # Update form fields with selected record
            self.name_var.set(self.selected_record[1])
            self.age_var.set(self.selected_record[2])
            
            # For salary, we need to handle the formatted display
            salary_value = self.selected_record[3]
            self.salary_var.set(salary_value)
            self.salary_entry.config(foreground="black")
            
            # Enable print button
            self.print_btn.config(state=NORMAL)
            
            # Update status
            self.status_label.config(text=f"Selected: {self.selected_record[1]} (ID: {self.selected_record[0]})")
    
    def clear_selection(self):
        """Clear selection and form fields"""
        self.tree.selection_remove(self.tree.selection())
        self.selected_record = None
        self.name_var.set("")
        self.age_var.set("")
        self.salary_var.set("")
        self.add_placeholder(None)  # Restore placeholder
        self.print_btn.config(state=DISABLED)
        self.status_label.config(text="No record selected")
    
    def submit_data(self):
        """Insert data into MySQL database"""
        name = self.name_var.get().strip()
        age = self.age_var.get().strip()
        salary_input = self.salary_var.get().strip()
        
        # Remove placeholder text if present
        if salary_input == "e.g., 50000 or 5.5 lakh":
            salary_input = ""
        
        # Validation
        if not name or not age or not salary_input:
            messagebox.showerror("Error", "Please fill all fields!")
            return
        
        try:
            age = int(age)
            salary = self.parse_currency(salary_input)
        except ValueError as e:
            messagebox.showerror("Error", str(e))
            return
        
        conn = self.create_connection()
        if conn:
            try:
                cursor = conn.cursor()
                query = "INSERT INTO Table1 (name, age, salary) VALUES (%s, %s, %s)"
                cursor.execute(query, (name, age, salary))
                conn.commit()
                
                messagebox.showinfo("Success", "Data inserted successfully!")
                
                # Clear form and selection
                self.clear_selection()
                
                # Refresh table
                self.load_data()
                
            except Error as e:
                messagebox.showerror("Database Error", f"Error inserting data: {e}")
            finally:
                cursor.close()
                conn.close()
    
    def load_data(self):
        """Load data from MySQL database and display in table"""
        # Clear existing data
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        conn = self.create_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT id, name, age, salary, created_at FROM Table1 ORDER BY created_at DESC")
                rows = cursor.fetchall()
                
                for row in rows:
                    # Format salary with Indian Rupees symbol
                    formatted_salary = self.format_currency(row[3])
                    
                    formatted_row = (
                        row[0],
                        row[1],
                        row[2],
                        formatted_salary,
                        row[4].strftime("%Y-%m-%d %H:%M") if row[4] else ""
                    )
                    self.tree.insert("", END, values=formatted_row)
                
            except Error as e:
                messagebox.showerror("Database Error", f"Error loading data: {e}")
            finally:
                cursor.close()
                conn.close()
    
    def populate_template(self, output_docx_path):
        """Populate the template with employee data"""
        try:
            # Check if template exists
            if not os.path.exists(self.template_path):
                messagebox.showerror("Error", f"Template file not found:\n{self.template_path}")
                return False
            
            # Load the template
            doc = Document(self.template_path)
            
            # Get raw salary value from database
            conn = self.create_connection()
            raw_salary = None
            if conn:
                try:
                    cursor = conn.cursor()
                    cursor.execute("SELECT salary FROM Table1 WHERE id = %s", (self.selected_record[0],))
                    result = cursor.fetchone()
                    if result:
                        raw_salary = result[0]
                    cursor.close()
                    conn.close()
                except Error:
                    raw_salary = self.parse_currency(self.selected_record[3])
            
            # Format salary
            formatted_salary = self.format_currency(raw_salary) if raw_salary is not None else self.selected_record[3]
            
            # Replacement data
            replacements = {
                '{{employee_id}}': str(self.selected_record[0]),
                '{{employee_name}}': self.selected_record[1],
                '{{employee_age}}': str(self.selected_record[2]),
                '{{employee_salary}}': formatted_salary,
                '{{generation_date}}': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # Replace placeholders in all paragraphs
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            
            # Replace placeholders in all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, value)
            
            # Replace placeholders in headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
                
                for paragraph in section.footer.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
            
            # Save the populated document
            doc.save(output_docx_path)
            return True
            
        except Exception as e:
            messagebox.showerror("Error", f"Error populating template: {e}")
            return False
    
    def convert_to_pdf(self, docx_path, pdf_path):
        """Convert DOCX to PDF using docx2pdf"""
        try:
            # Initialize COM for Windows
            if sys.platform.startswith('win'):
                pythoncom.CoInitialize()
            
            # Convert using docx2pdf
            convert(docx_path, pdf_path)
            return True
            
        except Exception as e:
            messagebox.showerror("Error", f"Error converting to PDF: {e}")
            return False
        finally:
            if sys.platform.startswith('win'):
                pythoncom.CoUninitialize()
    
    def generate_pdf(self):
        """Generate PDF from template for selected record"""
        if not self.selected_record:
            messagebox.showerror("Error", "No record selected!")
            return
        
        try:
            # Create output directory
            temp_dir = tempfile.gettempdir()
            output_dir = os.path.join(temp_dir, "EmployeeRecords")
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filenames
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename_base = f"Employee_{self.selected_record[0]}_{self.selected_record[1].replace(' ', '_')}"
            
            # Temporary DOCX file (will be deleted after conversion)
            temp_docx_path = os.path.join(output_dir, f"{filename_base}_{timestamp}.docx")
            pdf_path = os.path.join(output_dir, f"{filename_base}_{timestamp}.pdf")
            
            # Populate template
            if not self.populate_template(temp_docx_path):
                return
            
            # Convert to PDF
            if not self.convert_to_pdf(temp_docx_path, pdf_path):
                return
            
            # Delete temporary DOCX file (optional - comment out if you want to keep it)
            try:
                os.remove(temp_docx_path)
            except:
                pass  # Ignore if file deletion fails
            
            # Show success message
            messagebox.showinfo("Success", 
                               f"PDF document generated successfully!\n\n"
                               f"PDF: {pdf_path}\n\n"
                               f"File saved in: {output_dir}")
            
            # Optionally open the PDF
            try:
                os.startfile(pdf_path)  # For Windows
            except:
                try:
                    subprocess.run(['open', pdf_path])  # For macOS
                except:
                    try:
                        subprocess.run(['xdg-open', pdf_path])  # For Linux
                    except:
                        pass  # Skip if can't open file
            
        except Exception as e:
            messagebox.showerror("Error", f"Error generating PDF: {e}")

def main():
    # Create the main window
    root = tb.Window(themename="cosmo")
    app = EmployeeFormApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()